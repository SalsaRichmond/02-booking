import os
import sys

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx is not installed in the active environment. Run: py -m pip install python-docx")
    docx = None

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_recommendations_doc(filepath):
    if not docx:
        print("Cannot create docx file: python-docx is required.")
        return
    doc = docx.Document()

    # Set normal margins (0.8 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Document Header / Banner
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("SALSA GUY RICHMOND LLC — SYSTEM AUDIT & RECOMMENDATIONS")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(220, 38, 38) # Salsa Crimson

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Strategic Improvements for Booking Questionnaire, Google Apps Script Pipeline & Master Google Docs")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(100, 116, 139) # Slate

    # Metadata callout box
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_cell = meta_table.cell(0, 0)
    set_cell_background(meta_cell, "F8FAFC")
    set_cell_margins(meta_cell, top=140, bottom=140, left=200, right=200)
    
    meta_p = meta_cell.paragraphs[0]
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(0)
    meta_p.paragraph_format.line_spacing = 1.25
    
    r = meta_p.add_run("Target Audience: ")
    r.font.bold = True
    meta_p.add_run("Angel A. Rodriguez (The Salsa Guy / Profe)\n")
    
    r = meta_p.add_run("Subject Codebase: ")
    r.font.bold = True
    meta_p.add_run("Event Booking 2026 (public/index.html, Code.gs v20.71, Google Drive)\n")
    
    r = meta_p.add_run("Current Live URL: ")
    r.font.bold = True
    meta_p.add_run("https://sgr-booking-2026.tradicion.workers.dev\n")

    r = meta_p.add_run("Audit Scope: ")
    r.font.bold = True
    meta_p.add_run("Frontend form flow, backend JSON field mapping (lines 1015–1080), header synchronization, template document population, and mobile conversion.")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Helper function for adding styled sections
    def add_section_header(num, title, badge_text, badge_color):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        
        r_num = p.add_run(f"RECOMMENDATION {num}: ")
        r_num.font.name = "Arial"
        r_num.font.size = Pt(13)
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(15, 23, 42)
        
        r_title = p.add_run(title)
        r_title.font.name = "Arial"
        r_title.font.size = Pt(13)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(220, 38, 38)
        
        r_badge = p.add_run(f"  [{badge_text}]")
        r_badge.font.name = "Arial"
        r_badge.font.size = Pt(9.5)
        r_badge.font.bold = True
        if badge_color == "RED":
            r_badge.font.color.rgb = RGBColor(220, 38, 38)
        elif badge_color == "AMBER":
            r_badge.font.color.rgb = RGBColor(217, 119, 6)
        elif badge_color == "BLUE":
            r_badge.font.color.rgb = RGBColor(37, 99, 235)
        elif badge_color == "GREEN":
            r_badge.font.color.rgb = RGBColor(16, 185, 129)

    # 1. FIELD ALIGNMENT & DATA LOSS
    add_section_header("1", "Fix Missing Backend Mappings in Code.gs (Critical Data Loss Prevention)", "PRIORITY: HIGH (BUG FIX)", "RED")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("When examining lines 1015 to 1080 of ").font.color.rgb = RGBColor(51, 65, 85)
    r = p.add_run("Code.gs (handleFormSubmitJson)")
    r.font.bold = True
    p.add_run(" alongside the frontend submission payload in ").font.color.rgb = RGBColor(51, 65, 85)
    r2 = p.add_run("public/index.html")
    r2.font.bold = True
    p.add_run(", we identified multiple fields where the customer enters valuable information on the website, but the backend either lacks a matching field map or maps to the wrong key:\n")

    items_1 = [
        ("Missing 'outOfStateLogistics' Mapping: ", "The frontend collects detailed out-of-state flight/hotel logistics in formData.outOfStateLogistics, but fieldHeaderMap in Code.gs does NOT contain an entry for outOfStateLogistics. When out-of-state clients submit their lodging and per diem requirements, this data is never saved into the Google Sheet columns."),
        ("Duplicate 'eventPurpose' vs 'eventDescription': ", "In public/index.html line 3763, eventPurpose is set to document.getElementById('eventDescription').value. Meanwhile, the actual event purpose field is bypassed. Code.gs expects eventPurpose and eventDescription separately."),
        ("Missing 'serviceTypeRequested' Mapping: ", "The customer picks 'Performance Only', 'Dance Instruction Only', or 'Both', yet fieldHeaderMap lacks an explicit alias mapping for serviceTypeRequested to save the primary service choice in its own dedicated sheet column."),
        ("Repertoire Individual Category Storage: ", "The form collects all 21 repertoire checkboxes into a single concatenated string in performanceServices. However, Code.gs setupMasterHeaders defines separate columns for Repertoire: Mexico, Repertoire: Caribbean, Repertoire: Central America, Repertoire: South America, and Repertoire: Theatrical. Splitting them into their regional categories during submission ensures the Google Sheet and Review Doc reflect clean categorical grouping.")
    ]
    for b_title, b_desc in items_1:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.font.bold = True
        r_bt.font.color.rgb = RGBColor(15, 23, 42)
        r_bd = bp.add_run(b_desc)
        r_bd.font.color.rgb = RGBColor(51, 65, 85)

    # 2. AUTO-SAVE & DRAFT RESUME
    add_section_header("2", "Add LocalStorage Auto-Save & Draft Recovery to Prevent Abandoned Submissions", "PRIORITY: HIGH (CONVERSION)", "RED")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The 2026 booking questionnaire is comprehensive (covering 6 thorough sections from logistics to hospitality). Prospective event hosts often need time to look up their venue address, budget confirmation, or schedule before completing the form.")
    
    items_2 = [
        ("Real-Time Input Caching: ", "Automatically save client inputs to browser localStorage on every keystroke or selection. If a client accidentally closes their mobile tab or navigates away to check venue details, their answers are instantly restored upon return."),
        ("Draft Recovery Banner: ", "Display a subtle notification: 'We restored your previous event draft. [Clear Draft]' so users know their progress was saved."),
        ("Submission Clearance: ", "Clear the local cache only after handleFormSubmitJson returns a confirmed Request ID.")
    ]
    for b_title, b_desc in items_2:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.font.bold = True
        r_bd = bp.add_run(b_desc)
        r_bd.font.color.rgb = RGBColor(51, 65, 85)

    # 3. INTERACTIVE CALENDAR CONFLICT CHECKING
    add_section_header("3", "Pre-Submission Date Availability Indicator (Connected to Google Calendar)", "PRIORITY: MEDIUM (USER EXPERIENCE)", "AMBER")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Currently, the website mentions the general availability rule (Fridays–Sundays for shows, Mondays–Thursdays for dance lessons). However, clients can still submit requests for dates where Salsa Guy already has a booked performance in InfoCalendar.")

    items_3 = [
        ("Real-time Availability Status: ", "Expose a lightweight Apps Script endpoint (doGet with action=checkAvailability&date=YYYY-MM-DD) that queries InfoCalendar without revealing private event details."),
        ("Instant Visual Feedback: ", "When a date is selected, show: '✅ Date Appears Available', '⚠️ Date Has Existing Bookings (Inquire for Alternative Hours)', or 'ℹ️ Weekday: Available for Instruction/Workshops Only'."),
        ("Reduces Back-and-Forth: ", "Saves hours of email correspondence by setting clear scheduling expectations right at the moment of questionnaire fill-in.")
    ]
    for b_title, b_desc in items_3:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.font.bold = True
        r_bd = bp.add_run(b_desc)
        r_bd.font.color.rgb = RGBColor(51, 65, 85)

    # 4. INSTANT AUTO-GENERATED QUOTE / ESTIMATE
    add_section_header("4", "Dynamic Price Range Calculator (Instant Budget Transparency)", "PRIORITY: MEDIUM (SALES VELOCITY)", "AMBER")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Event organizers frequently hesitate when selecting 'Confirmed Budget Amount' because they do not know standard rates for Latin dance performances, live Parranda caroling, or sound system provisioning.")

    items_4 = [
        ("Dynamic Estimate Widget: ", "Based on service type (Solo instruction vs. 4-person troupe vs. full Parranda ensemble) and duration (1 hour vs. half-day), provide an estimated ballpark investment range right above the budget field."),
        ("Filters Unqualified Leads: ", "Educates clients on professional performance rates before they submit unrealistic numbers (e.g. $100 for a 6-performer Mexican folkloric production)."),
        ("Custom Package Tiering: ", "Add quick-select tier buttons: 'Standard Cultural Showcase', 'Festival Headliner Experience', or 'VIP Masterclass + Performance'.")
    ]
    for b_title, b_desc in items_4:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.font.bold = True
        r_bd = bp.add_run(b_desc)
        r_bd.font.color.rgb = RGBColor(51, 65, 85)

    # 5. WHATSAPP & SMS INSTANT NOTIFICATIONS
    add_section_header("5", "Automated SMS / WhatsApp Notification for Profe & Client", "PRIORITY: MEDIUM (RESPONSE SPEED)", "BLUE")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("Currently, Code.gs sends automated HTML emails via sendAdminSubmittalNotification and sendClientReceiptNotification. While email works well, booking inquiries are time-sensitive.")

    items_5 = [
        ("Instant SMS / WhatsApp Ping for Angel: ", "Integrate a webhook (e.g., Twilio or WhatsApp Business API via Google Apps Script UrlFetchApp) sending a 1-sentence notification: '💃 New Booking Request! Maria Santos on Oct 14 for $1,500 at Richmond Center. Tap to view Proposal Doc.'"),
        ("Fast 15-Minute Response Time: ", "Event planners often book the first vendor that responds. Rapid notification gives Salsa Guy a massive competitive advantage.")
    ]
    for b_title, b_desc in items_5:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.font.bold = True
        r_bd = bp.add_run(b_desc)
        r_bd.font.color.rgb = RGBColor(51, 65, 85)

    # 6. MASTER GOOGLE DOC TEMPLATE OPTIMIZATION
    add_section_header("6", "Master Google Docs Template Modernization (Proposal, Contract, Perf Info)", "PRIORITY: HIGH (BRANDING & POLISH)", "RED")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The automation suite generates three customized Google Docs in Drive for every lead using templates: Proposal, Contract, and Performance Information. Our inspection of Code.gs lines 1980–2015 revealed opportunities to elevate these documents:")

    items_6 = [
        ("Automated Dynamic Repertoire Table: ", "Instead of dumping raw checkbox text, generate a styled, formatted table in the Proposal and Performance Info Doc listing each selected dance (e.g., 'El Baile de los Viejitos', 'Bomba', 'Alma Ranchera') along with performer count and required stage dimensions."),
        ("Google Doc PDF Auto-Export & Direct Download Link: ", "In addition to generating the editable Google Doc, have processRow generate a finalized read-only PDF using doc.getAs('application/pdf') and save it to the event's Drive folder. This link can be immediately sent to the client in their receipt email."),
        ("Digital Signature Readiness: ", "Add an interactive signing block placeholder with date and IP stamp to prepare the contract document for e-signature workflows.")
    ]
    for b_title, b_desc in items_6:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r_bt = bp.add_run(b_title)
        r_bt.font.bold = True
        r_bd = bp.add_run(b_desc)
        r_bd.font.color.rgb = RGBColor(51, 65, 85)

    # 7. MULTI-LANGUAGE DEEP LOCALIZATION
    add_section_header("7", "Complete Multi-Language Field Synchronization (Spanish, French, German, etc.)", "PRIORITY: MEDIUM (CULTURAL & REACH)", "BLUE")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.add_run("The frontend has an impressive language switcher supporting English, Spanish, French, German, etc. However, following the reorganization into 6 sequential sections, some translation strings in langDict still reference legacy section numbers (e.g., 'sec6: 7. Technical...', 'sec5: 8. Budget...'). Synchronizing these dictionary keys guarantees seamless multi-lingual presentation for international and cultural embassy clients.")

    # 8. SUMMARY MATRIX TABLE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r_mat = p.add_run("EXECUTIVE SUMMARY & ACTION ROADMAP")
    r_mat.font.name = "Arial"
    r_mat.font.size = Pt(14)
    r_mat.font.bold = True
    r_mat.font.color.rgb = RGBColor(15, 23, 42)

    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(0.6), Inches(2.8), Inches(1.8), Inches(1.5)]
    headers_text = ["#", "Recommendation", "Impact Area", "Target Timeline"]

    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers_text):
        hdr_cells[i].text = h_text
        set_cell_background(hdr_cells[i], "DC2626")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)

    roadmap_data = [
        ("1", "Fix Code.gs Field Mappings & Aliases", "Data Integrity & Sheets Logging", "Phase 1 (Immediate)"),
        ("2", "LocalStorage Draft Auto-Save & Recovery", "User Experience & Lead Recovery", "Phase 1 (Immediate)"),
        ("3", "Calendar Availability Real-Time Check", "Client Scheduling Accuracy", "Phase 2 (Next Sprint)"),
        ("4", "Dynamic Price Range Ballpark Widget", "Lead Qualification & Conversion", "Phase 2 (Next Sprint)"),
        ("5", "SMS / WhatsApp Instant Host Notifications", "Lead Response Time (< 15 mins)", "Phase 3 (Expansion)"),
        ("6", "Master Doc PDF Auto-Export & Styled Tables", "Proposal Presentation & Close Rate", "Phase 2 (Next Sprint)"),
        ("7", "Complete 6-Section i18n Translation Sync", "Global & Cultural Client Reach", "Phase 2 (Next Sprint)")
    ]

    for row_idx, data in enumerate(roadmap_data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 65, 85)

    # Set column widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Conclusion & Signoff
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    end_p = doc.add_paragraph()
    end_p.paragraph_format.space_before = Pt(10)
    end_p.paragraph_format.line_spacing = 1.2
    r_end = end_p.add_run("Prepared with dedication to the continued excellence and growth of Salsa Guy Richmond LLC & Tradición Puerto Rican Folk Dancing.\n")
    r_end.font.italic = True
    r_end.font.color.rgb = RGBColor(100, 116, 139)
    r_sign = end_p.add_run("Ready for immediate implementation upon review.")
    r_sign.font.bold = True
    r_sign.font.color.rgb = RGBColor(220, 38, 38)

    doc.save(filepath)
    print(f"Document saved successfully to {filepath}")

if __name__ == "__main__":
    out_dir = r"G:\My Drive"
    out_path = os.path.join(out_dir, "Read this Salsa Guy.docx")
    create_recommendations_doc(out_path)
